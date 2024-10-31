import os
import re
import pandas as pd
import streamlit as st
from io import BytesIO
from docx import Document
import PyPDF2

def extract_emails_from_files(files, log_callback=None):
    email_set = set()
    # Enhanced regex pattern with word boundaries
    email_pattern = re.compile(
        r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b', re.IGNORECASE
    )
    
    total_files = len(files)
    
    for idx, uploaded_file in enumerate(files, start=1):
        filename = uploaded_file.name
        file_extension = os.path.splitext(filename)[1].lower()
        
        if log_callback:
            log_callback(f"Processing file {idx}/{total_files}: {filename}")
        
        try:
            if file_extension in ['.xls', '.xlsx', '.xlsm']:
                # Process Excel files
                xls = pd.ExcelFile(uploaded_file)
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str)
                    for column in df.columns:
                        for cell in df[column].dropna():
                            matches = email_pattern.findall(str(cell))
                            email_set.update(matches)
            elif file_extension == '.docx':
                # Process Word .docx documents
                doc = Document(uploaded_file)
                # Extract text from paragraphs
                for para in doc.paragraphs:
                    matches = email_pattern.findall(para.text)
                    email_set.update(matches)
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            matches = email_pattern.findall(cell.text)
                            email_set.update(matches)
            elif file_extension == '.pdf':
                # Process PDF files
                reader = PyPDF2.PdfReader(uploaded_file)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text = page.extract_text()
                    if text:
                        matches = email_pattern.findall(text)
                        email_set.update(matches)
            else:
                if log_callback:
                    log_callback(f"Unsupported file type: {filename}")
        except Exception as e:
            if log_callback:
                log_callback(f"Error processing file {filename}: {e}")
    
    return sorted(email_set)

def convert_df_to_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Email Addresses')
    return output.getvalue()

def convert_df_to_csv(df):
    return df.to_csv(index=False).encode('utf-8')

def convert_df_to_txt(df):
    return '\n'.join(df['Email Addresses'].tolist()).encode('utf-8')

def main():
    st.set_page_config(
        page_title="📧 Email Address Extractor",
        page_icon="📧",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    
    st.title("📧 Email Address Extractor from Various File Types")
    st.write("""
    Upload multiple files (`.xls`, `.xlsx`, `.xlsm`, `.docx`, `.pdf`) containing email addresses in various formats.
    The app will extract all unique email addresses and provide a consolidated file for download.
    
    **Note:** Scanned PDFs (containing images of text) are not supported. Please ensure your PDFs are text-based or use an OCR tool to convert them before uploading.
    **Note:** To handle `.doc` files, please convert them to `.docx` before uploading.
    """)
    
    # Instructions
    st.markdown("""
    ### Instructions:
    1. **Upload Files:** Click on the "Choose Files" button to upload multiple Excel (`.xls`, `.xlsx`, `.xlsm`), Word (`.docx`), or PDF (`.pdf`) files.
    2. **Select Download Format:** Choose your preferred format for the extracted emails.
    3. **Filter Emails:** Use the filter box to search for specific email addresses.
    4. **Extract Emails:** Click the "Extract Emails" button to start the process.
    5. **Download Results:** Once extraction is complete, download the results in your chosen format.
    """)
    
    # File Uploader and Download Format Selection in Columns
    col1, col2 = st.columns([3, 1])
    with col1:
        uploaded_files = st.file_uploader(
            "Choose Files",
            type=['xls', 'xlsx', 'xlsm', 'docx', 'pdf'],
            accept_multiple_files=True
        )
    with col2:
        download_format = st.selectbox(
            "Select Download Format",
            options=["Excel (.xlsx)", "CSV (.csv)", "Text (.txt)"]
        )
    
    # Search/Filter Box
    filter_text = st.text_input("Filter Emails", help="Search for specific email addresses.")
    
    # Initialize session state for logs
    if 'logs' not in st.session_state:
        st.session_state.logs = []
    
    # Function to log messages
    def log(message):
        st.session_state.logs.append(message)
        # Scroll to the bottom (Streamlit doesn't support auto-scrolling, so rerun is used)
        st.experimental_rerun()
    
    # Extraction Button
    if st.button("Extract Emails"):
        if uploaded_files:
            st.session_state.logs = []  # Reset logs
            with st.spinner('Processing...'):
                # Display a progress bar
                progress_bar = st.progress(0)
                total_files = len(uploaded_files)
                
                def log_callback(message):
                    log(message)
                    # Update progress
                    progress = len(st.session_state.logs) / (total_files * 2)  # Approximation
                    progress_bar.progress(min(progress, 1.0))
                
                emails = extract_emails_from_files(uploaded_files, log_callback=log_callback)
                progress_bar.progress(1.0)
            
            if emails:
                st.success(f"Extraction complete! {len(emails)} unique email(s) found.")
                
                # Create DataFrame
                email_df = pd.DataFrame({'Email Addresses': emails})
                
                # Apply filter if any
                if filter_text:
                    filtered_df = email_df[email_df['Email Addresses'].str.contains(filter_text, case=False, na=False)]
                else:
                    filtered_df = email_df
                
                # Display the DataFrame with filtering
                st.subheader("Extracted Emails")
                st.dataframe(filtered_df)
                
                # Download Buttons
                if download_format == "Excel (.xlsx)":
                    excel_data = convert_df_to_excel(filtered_df)
                    st.download_button(
                        label="📥 Download as Excel",
                        data=excel_data,
                        file_name='Extracted_Email_Addresses.xlsx',
                        mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                    )
                elif download_format == "CSV (.csv)":
                    csv_data = convert_df_to_csv(filtered_df)
                    st.download_button(
                        label="📥 Download as CSV",
                        data=csv_data,
                        file_name='Extracted_Email_Addresses.csv',
                        mime='text/csv'
                    )
                elif download_format == "Text (.txt)":
                    txt_data = convert_df_to_txt(filtered_df)
                    st.download_button(
                        label="📥 Download as TXT",
                        data=txt_data,
                        file_name='Extracted_Email_Addresses.txt',
                        mime='text/plain'
                    )
            else:
                st.warning("No email addresses found in the uploaded files.")
        else:
            st.warning("Please upload at least one file.")
    
    # Display Logs
    if st.session_state.logs:
        st.subheader("Logs")
        log_container = st.container()
        for log_msg in st.session_state.logs:
            log_container.write(f"- {log_msg}")
